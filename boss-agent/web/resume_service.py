"""
简历服务层 — 封装简历 CRUD 逻辑，供 API 路由调用。

职责：
- get_active_resume(): 读取活跃简历 + 求职意向
- update_resume(data): 白名单校验 → 写入 resumes / job_preferences
"""

from __future__ import annotations

import io
import json
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from db.database import Database

SCALAR_FIELDS = {
    "name", "phone", "email", "birth_year", "city",
    "education_level", "school", "education_major",
    "years_of_experience", "current_company", "current_role",
    "summary", "self_evaluation", "raw_text",
}

LIST_FIELDS = {
    "work_experience", "projects", "highlights", "skills_flat", "tech_stack",
}

_JOB_PREF_SCALAR = {"salary_min", "salary_max", "work_type"}
_JOB_PREF_LIST = {
    "target_cities", "target_roles", "priorities", "deal_breakers",
}


def _safe_json_load(val: str | None) -> Any:
    if not val:
        return None
    try:
        return json.loads(val)
    except (json.JSONDecodeError, TypeError):
        return val


class ResumeService:
    """简历服务层：CRUD + DOCX 导出"""

    def __init__(self, db: Database) -> None:
        self.db = db

    # ------------------------------------------------------------------
    # 读取
    # ------------------------------------------------------------------

    async def get_active_resume(self) -> dict | None:
        """获取当前活跃简历的完整数据（含 job_preferences）。"""
        rows = await self.db.execute(
            "SELECT * FROM resumes WHERE is_active = 1 ORDER BY id DESC LIMIT 1"
        )
        if not rows:
            return None

        resume = rows[0]
        result: dict[str, Any] = {}

        # 标量字段
        for field in SCALAR_FIELDS:
            result[field] = resume.get(field)

        # 列表字段 — JSON 反序列化
        for field in LIST_FIELDS:
            result[field] = _safe_json_load(resume.get(field))

        result["id"] = resume["id"]
        result["updated_at"] = resume.get("updated_at")

        # 关联 job_preferences
        pref_rows = await self.db.execute(
            "SELECT * FROM job_preferences WHERE resume_id = ? AND is_active = 1 LIMIT 1",
            (resume["id"],),
        )
        if pref_rows:
            pref = pref_rows[0]
            result["job_preferences"] = {
                "target_cities": _safe_json_load(pref.get("target_cities")),
                "target_roles": _safe_json_load(pref.get("target_roles")),
                "salary_min": pref.get("salary_min"),
                "salary_max": pref.get("salary_max"),
                "work_type": pref.get("work_type"),
                "priorities": _safe_json_load(pref.get("priorities")),
                "deal_breakers": _safe_json_load(pref.get("deal_breakers")),
            }
        else:
            result["job_preferences"] = None

        return result

    # ------------------------------------------------------------------
    # 更新
    # ------------------------------------------------------------------

    async def update_resume(self, data: dict) -> dict:
        """
        更新简历数据。

        - 白名单校验：忽略不在 SCALAR_FIELDS / LIST_FIELDS 中的字段
        - 标量字段直接更新 resumes 表
        - 列表字段 JSON 序列化后更新 resumes 表
        - job_preferences 子对象更新/创建 job_preferences 表
        - 无活跃简历时自动创建
        - 每次更新刷新 updated_at

        返回 {"updated": True, "fields": [已更新字段名列表]}
        """
        # 获取或创建活跃简历
        rows = await self.db.execute(
            "SELECT id FROM resumes WHERE is_active = 1 ORDER BY id DESC LIMIT 1"
        )
        if rows:
            resume_id = rows[0]["id"]
        else:
            resume_id = await self.db.execute_write(
                "INSERT INTO resumes (is_active) VALUES (1)"
            )

        updated_fields: list[str] = []
        resume_updates: dict[str, Any] = {}

        # 标量字段
        for field in SCALAR_FIELDS:
            if field in data:
                resume_updates[field] = data[field]
                updated_fields.append(field)

        # 列表字段 — JSON 序列化
        for field in LIST_FIELDS:
            if field in data:
                resume_updates[field] = json.dumps(data[field], ensure_ascii=False)
                updated_fields.append(field)

        # 写入 resumes 表
        if resume_updates:
            set_parts = [f"{k} = ?" for k in resume_updates]
            set_parts.append("updated_at = CURRENT_TIMESTAMP")
            set_clause = ", ".join(set_parts)
            values = list(resume_updates.values()) + [resume_id]
            await self.db.execute_write(
                f"UPDATE resumes SET {set_clause} WHERE id = ?",
                tuple(values),
            )
        else:
            # 即使没有简历字段更新，也刷新 updated_at
            await self.db.execute_write(
                "UPDATE resumes SET updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                (resume_id,),
            )

        # 处理 job_preferences
        job_prefs = data.get("job_preferences")
        if isinstance(job_prefs, dict):
            pref_updates: dict[str, Any] = {}
            for field in _JOB_PREF_SCALAR:
                if field in job_prefs:
                    pref_updates[field] = job_prefs[field]
            for field in _JOB_PREF_LIST:
                if field in job_prefs:
                    val = job_prefs[field]
                    if isinstance(val, (list, dict)):
                        val = json.dumps(val, ensure_ascii=False)
                    pref_updates[field] = val

            if pref_updates:
                pref_rows = await self.db.execute(
                    "SELECT id FROM job_preferences WHERE resume_id = ? AND is_active = 1 LIMIT 1",
                    (resume_id,),
                )
                if pref_rows:
                    pref_id = pref_rows[0]["id"]
                    set_clause = ", ".join(f"{k} = ?" for k in pref_updates)
                    values = list(pref_updates.values()) + [pref_id]
                    await self.db.execute_write(
                        f"UPDATE job_preferences SET {set_clause}, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                        tuple(values),
                    )
                else:
                    pref_updates["resume_id"] = resume_id
                    cols = ", ".join(pref_updates.keys())
                    placeholders = ", ".join("?" for _ in pref_updates)
                    await self.db.execute_write(
                        f"INSERT INTO job_preferences ({cols}) VALUES ({placeholders})",
                        tuple(pref_updates.values()),
                    )
                updated_fields.append("job_preferences")

        return {"updated": True, "fields": updated_fields}

    # ------------------------------------------------------------------
    # DOCX 导出
    # ------------------------------------------------------------------

    async def export_docx(self) -> tuple[bytes, str]:
        """
        导出当前活跃简历为 DOCX。

        返回 (docx_bytes, filename)
        filename 格式: "简历_{name}_{YYYYMMDD}.docx"
        无活跃简历时抛出 ValueError
        """
        resume = await self.get_active_resume()
        if not resume:
            raise ValueError("暂无简历数据")

        doc = Document()

        # --- 标题：姓名 ---
        name = resume.get("name") or "未命名"
        title_para = doc.add_heading(name, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- 联系方式 ---
        contact_parts: list[str] = []
        if resume.get("phone"):
            contact_parts.append(resume["phone"])
        if resume.get("email"):
            contact_parts.append(resume["email"])
        if resume.get("city"):
            contact_parts.append(resume["city"])
        if contact_parts:
            p = doc.add_paragraph(" | ".join(contact_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- 个人简介 ---
        if resume.get("summary"):
            doc.add_heading("个人简介", level=1)
            doc.add_paragraph(resume["summary"])

        # --- 自我评价 ---
        if resume.get("self_evaluation"):
            doc.add_heading("自我评价", level=1)
            doc.add_paragraph(resume["self_evaluation"])

        # --- 技术栈 ---
        tech_stack = resume.get("tech_stack")
        if tech_stack:
            doc.add_heading("技术栈", level=1)
            if isinstance(tech_stack, dict):
                for category, skills in tech_stack.items():
                    if isinstance(skills, list):
                        doc.add_paragraph(f"{category}: {', '.join(str(s) for s in skills)}")
                    else:
                        doc.add_paragraph(f"{category}: {skills}")
            elif isinstance(tech_stack, list):
                doc.add_paragraph(", ".join(str(s) for s in tech_stack))

        # --- 工作经历 ---
        work_exp = resume.get("work_experience")
        if work_exp and isinstance(work_exp, list):
            doc.add_heading("工作经历", level=1)
            for entry in work_exp:
                if not isinstance(entry, dict):
                    continue
                company = entry.get("company", "")
                role = entry.get("role", "")
                duration = entry.get("duration", "")
                header = f"{company} · {role}"
                if duration:
                    header += f" ({duration})"
                doc.add_heading(header, level=2)
                if entry.get("description"):
                    doc.add_paragraph(entry["description"])
                highlights = entry.get("highlights")
                if highlights and isinstance(highlights, list):
                    for h in highlights:
                        doc.add_paragraph(str(h), style="List Bullet")

        # --- 项目经验 ---
        projects = resume.get("projects")
        if projects and isinstance(projects, list):
            doc.add_heading("项目经验", level=1)
            for proj in projects:
                if not isinstance(proj, dict):
                    continue
                proj_name = proj.get("name", "")
                doc.add_heading(proj_name, level=2)
                if proj.get("description"):
                    doc.add_paragraph(proj["description"])
                highlights = proj.get("highlights")
                if highlights and isinstance(highlights, list):
                    for h in highlights:
                        doc.add_paragraph(str(h), style="List Bullet")

        # --- 核心亮点 ---
        highlights_top = resume.get("highlights")
        if highlights_top and isinstance(highlights_top, list):
            doc.add_heading("核心亮点", level=1)
            for h in highlights_top:
                doc.add_paragraph(str(h), style="List Bullet")

        # --- 求职意向 ---
        prefs = resume.get("job_preferences")
        if prefs and isinstance(prefs, dict):
            doc.add_heading("求职意向", level=1)
            if prefs.get("target_cities"):
                cities = prefs["target_cities"]
                if isinstance(cities, list):
                    doc.add_paragraph(f"目标城市: {', '.join(str(c) for c in cities)}")
            if prefs.get("target_roles"):
                roles = prefs["target_roles"]
                if isinstance(roles, list):
                    doc.add_paragraph(f"目标岗位: {', '.join(str(r) for r in roles)}")
            sal_parts: list[str] = []
            if prefs.get("salary_min") is not None:
                sal_parts.append(str(prefs["salary_min"]))
            if prefs.get("salary_max") is not None:
                sal_parts.append(str(prefs["salary_max"]))
            if sal_parts:
                doc.add_paragraph(f"期望薪资: {'-'.join(sal_parts)}K")
            if prefs.get("work_type"):
                doc.add_paragraph(f"工作类型: {prefs['work_type']}")
            if prefs.get("priorities") and isinstance(prefs["priorities"], list):
                doc.add_paragraph(f"优先考虑: {', '.join(str(p) for p in prefs['priorities'])}")
            if prefs.get("deal_breakers") and isinstance(prefs["deal_breakers"], list):
                doc.add_paragraph(f"不能接受: {', '.join(str(d) for d in prefs['deal_breakers'])}")

        # --- 序列化为 bytes ---
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        # --- 文件名 ---
        date_str = datetime.now().strftime("%Y%m%d")
        filename = f"简历_{name}_{date_str}.docx"

        return docx_bytes, filename
